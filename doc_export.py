"""G12: geracao dos documentos exportados (DOCX via python-docx, HTML
autocontido para /api/export-html -- G20) + os helpers build_*_rows
compartilhados pelos dois formatos e mask_secrets/doc_value (G7). Extraido
de backend.py na modularizacao G12.

Possui SEU PROPRIO lazy-loading de python-docx (load_docx() + Document/
DOCX_IMPORT_ERROR/WD_CELL_VERTICAL_ALIGNMENT/WD_TABLE_ALIGNMENT/
WD_ALIGN_PARAGRAPH/OxmlElement/qn/Inches/Pt/RGBColor) -- co-locado aqui
porque as dezenas de funcoes add_*/set_*/build_documentation_docx abaixo
referenciam esses nomes como bare (sem qualificar), exatamente como faziam
quando viviam todas juntas em backend.py; mover so build_documentation_docx
sem mover load_docx() junto quebraria toda essa cadeia de nomes bare. Esse
estado NAO e duplicado em nenhum outro modulo (single source of truth) --
render_graphics.py tem seu proprio estado independente para Pillow/cairosvg,
sem overlap.

`build_documentation_docx`/`build_documentation_html` tambem chamam
`PBIXRay(...)`/`load_pbixray()` diretamente (alem de analyze_pbix) -- por
isso, IGUAL a pbix_analysis.py, usam `backend.PBIXRay(...)`/
`backend.load_pbixray()`/`backend.PBIXRAY_IMPORT_ERROR` qualificados (ver o
docstring de pbix_analysis.py para o motivo completo: o monkeypatch dos
testes so alcanca o namespace de `backend`). Pelo MESMO motivo que
pbix_analysis.py, `import backend` e feito DENTRO do corpo de cada uma
dessas duas funcoes (nao no topo do modulo) -- evita que um
`import doc_export` standalone (sem `backend` ja carregado) dispare o mesmo
ciclo parcialmente-inicializado.

`add_cover()` precisa do estado PIL de render_graphics.py (para o banner
dourado da capa) -- acessa via `import render_graphics` +
`render_graphics.PIL_AVAILABLE`/`render_graphics._PILImage`/
`render_graphics._PILDraw` qualificados, pelo MESMO motivo (evitar copiar
por valor um estado que so existe de verdade depois de load_pillow() rodar).
"""
from __future__ import annotations

import base64
import html
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

import render_graphics
from bi_server import ROOT
from graph_utils import doc_value, first_value, mask_secrets, record_text
from logging_setup import logger
from pbix_analysis import analyze_pbix, is_system_table, records_from
from render_graphics import (
    _pil_png,
    build_architecture_svg,
    build_erd_svg,
    load_pillow,
    make_arch_png,
    make_banner_png,
    make_erd_png,
    svg_to_png_bytes,
)


Document = None
DOCX_IMPORT_ERROR = ""
_DOCX_IMPORT_ATTEMPTED = False


def load_docx() -> bool:
    global Document, DOCX_IMPORT_ERROR, _DOCX_IMPORT_ATTEMPTED
    global WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH
    global OxmlElement, qn, Inches, Pt, RGBColor
    if not _DOCX_IMPORT_ATTEMPTED:
        _DOCX_IMPORT_ATTEMPTED = True
        try:
            from docx import Document as _Document
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as _WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement as _OxmlElement
            from docx.oxml.ns import qn as _qn
            from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGBColor
            Document = _Document
            WD_CELL_VERTICAL_ALIGNMENT = _WD_CELL_VERTICAL_ALIGNMENT
            WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
            WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
            OxmlElement = _OxmlElement
            qn = _qn
            Inches, Pt, RGBColor = _Inches, _Pt, _RGBColor
        except Exception as import_error:
            DOCX_IMPORT_ERROR = str(import_error)
            logger.warning("python-docx nao pode ser importado: %s", import_error, exc_info=True)
    return Document is not None


def insert_svg_image(doc, svg_string: str, width_inches: float = 6.5, caption: str = "") -> bool:
    """Render SVG to PNG and insert into the document. Returns True on success."""
    if not svg_string:
        return False
    png_bytes = svg_to_png_bytes(svg_string, scale=2.0)
    if not png_bytes:
        return False
    buf = BytesIO(png_bytes)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run().add_picture(buf, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cap_run = cap.add_run(caption)
        cap_run.font.size = Pt(8.5)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor(96, 94, 92)
    return True


DOC_TEXT = {
    "pt-BR": {
        "metrics": {
            "sources": "Fontes",
            "queries": "Queries",
            "tables": "Tabelas",
            "measures": "Medidas",
            "calc_columns": "Colunas calculadas",
            "relationships": "Relacionamentos",
            "pages": "Páginas",
            "visuals": "Visuais",
        },
        "sections": [
            "Visão geral",
            "Fontes de dados",
            "Power Query",
            "Modelo semântico",
            "Dicionário de dados",
            "Medidas DAX",
            "Colunas calculadas",
            "Relacionamentos",
            "Páginas e visuais",
            "Linhagem técnica",
            "Diagnósticos",
        ],
        "cover_subtitle": "Documentação técnica do painel",
        "generated_at": "Gerado em",
        "toc": "Sumário",
        "overview_banner": "Visão geral do relatório",
        "overview_text": "Documentação gerada automaticamente pelo BI Flow Mapper a partir dos metadados locais do arquivo PBIX.",
        "file": "Arquivo",
        "analysis_source": "Origem da análise",
        "total_nodes": "Total de nós no mapa",
        "total_edges": "Total de arestas de linhagem",
        "data_sources_banner": "Conectores e origens de dados detectados",
        "source_headers": ["Fonte", "Padrão detectado", "Caminho / servidor"],
        "no_sources": "Nenhuma fonte detectada.",
        "arch_caption": "Diagrama de arquitetura — fontes conectadas ao dataset Power BI",
        "sources_detected": "fonte(s) detectada(s):",
        "connection_details": "2.1 Detalhes de conexão (TMSCHEMA_DATASOURCES)",
        "connection_headers": ["Tipo", "Nome", "String de conexão"],
        "power_query_banner": "Transformações e consultas M",
        "query_headers": ["Query", "Conexão / caminho", "Resumo técnico"],
        "no_queries": "Nenhuma query encontrada.",
        "semantic_model_banner": "Tabelas carregadas no modelo",
        "table_headers": ["Tabela", "Colunas", "Oculta", "Descrição"],
        "no_tables": "Nenhuma tabela encontrada.",
        "dictionary_banner": "Colunas, tipos e metadados de cada tabela",
        "column_headers": ["Tabela", "Coluna", "Tipo", "Formato", "Oculta", "Expressão"],
        "no_columns": "Nenhuma coluna encontrada nos metadados.",
        "measures_banner": "Medidas calculadas com expressões DAX",
        "measure_headers": ["Tabela", "Medida"],
        "no_measures": "Nenhuma medida encontrada.",
        "calc_columns_banner": "Colunas derivadas por expressões DAX",
        "calc_headers": ["Tabela", "Coluna"],
        "no_calc_columns": "Nenhuma coluna calculada encontrada.",
        "relationships_banner": "Vínculos entre tabelas do modelo semântico",
        "relationship_headers": ["Tabela origem", "Coluna origem", "Tabela destino", "Coluna destino", "Card.", "Filtro cruzado", "Ativo"],
        "no_relationships": "Nenhum relacionamento encontrado.",
        "erd_caption": "Diagrama entidade-relacionamento — modelo semântico Power BI",
        "pages_banner": "Páginas do relatório e seus elementos visuais",
        "page_headers": ["#", "Página", "Visuais", "Canvas (px)"],
        "no_pages": "Nenhuma página encontrada.",
        "visual_headers": ["Visual", "Tipo", "Campos / referências"],
        "no_visuals": "Nenhum visual encontrado.",
        "lineage_banner": "Grafo completo de dependências do pipeline",
        "lineage_headers": ["Origem", "Relacao", "Destino"],
        "no_lineage": "Nenhuma aresta de linhagem encontrada.",
        "diagnostics_banner": "Alertas e informações técnicas da análise",
        "relationship_columns": "Colunas disponíveis na tabela de relacionamentos:",
        "rel_from": "Origem",
        "rel_to": "Destino",
        "cardinality": "Cardinalidade",
        "cross_filter": "Filtro cruzado",
        "relationship_example": "Exemplo de relacionamento detectado — ",
        "message": "Mensagem",
        "no_diagnostics": "Nenhum diagnóstico registrado.",
        "yes": "Sim",
        "no": "Não",
        "dash": "—",
        "m_code_heading": "3.1 Código M por consulta",
        "summary": "Resumo",
        "technical_expressions": "Expressões técnicas",
        "additional_records_omitted": "{count} registros adicionais foram omitidos nesta seção para manter o documento legível.",
        "no_records": "Nenhum registro encontrado.",
        "record": "Registro",
        "truncated": "truncado",
        "connector": "Conector",
        "steps": "Etapas",
        "functions": "Funções",
        "datatype_map": {
            "2": "texto", "3": "decimal", "4": "inteiro", "5": "decimal",
            "6": "moeda", "7": "data", "8": "booleano", "9": "binário",
            "10": "variant", "17": "inteiro",
            "string": "texto", "int64": "inteiro", "double": "decimal",
            "boolean": "booleano", "datetime": "data/hora", "binary": "binário",
        },
    },
    "en-US": {
        "metrics": {
            "sources": "Sources",
            "queries": "Queries",
            "tables": "Tables",
            "measures": "Measures",
            "calc_columns": "Calculated columns",
            "relationships": "Relationships",
            "pages": "Pages",
            "visuals": "Visuals",
        },
        "sections": [
            "Overview",
            "Data sources",
            "Power Query",
            "Semantic model",
            "Data dictionary",
            "DAX measures",
            "Calculated columns",
            "Relationships",
            "Pages and visuals",
            "Technical lineage",
            "Diagnostics",
        ],
        "cover_subtitle": "Technical report documentation",
        "generated_at": "Generated at",
        "toc": "Table of contents",
        "overview_banner": "Report overview",
        "overview_text": "Documentation automatically generated by BI Flow Mapper from the local metadata in the PBIX file.",
        "file": "File",
        "analysis_source": "Analysis source",
        "total_nodes": "Total nodes in the map",
        "total_edges": "Total lineage edges",
        "data_sources_banner": "Detected connectors and data sources",
        "source_headers": ["Source", "Detected pattern", "Path / server"],
        "no_sources": "No sources detected.",
        "arch_caption": "Architecture diagram — sources connected to the Power BI dataset",
        "sources_detected": "source(s) detected:",
        "connection_details": "2.1 Connection details (TMSCHEMA_DATASOURCES)",
        "connection_headers": ["Type", "Name", "Connection string"],
        "power_query_banner": "M transformations and queries",
        "query_headers": ["Query", "Connection / path", "Technical summary"],
        "no_queries": "No queries found.",
        "semantic_model_banner": "Tables loaded into the model",
        "table_headers": ["Table", "Columns", "Hidden", "Description"],
        "no_tables": "No tables found.",
        "dictionary_banner": "Columns, types, and metadata for each table",
        "column_headers": ["Table", "Column", "Type", "Format", "Hidden", "Expression"],
        "no_columns": "No columns found in the metadata.",
        "measures_banner": "Calculated measures with DAX expressions",
        "measure_headers": ["Table", "Measure"],
        "no_measures": "No measures found.",
        "calc_columns_banner": "Columns derived from DAX expressions",
        "calc_headers": ["Table", "Column"],
        "no_calc_columns": "No calculated columns found.",
        "relationships_banner": "Links between semantic model tables",
        "relationship_headers": ["Source table", "Source column", "Target table", "Target column", "Card.", "Cross filter", "Active"],
        "no_relationships": "No relationships found.",
        "erd_caption": "Entity relationship diagram — Power BI semantic model",
        "pages_banner": "Report pages and their visual elements",
        "page_headers": ["#", "Page", "Visuals", "Canvas (px)"],
        "no_pages": "No pages found.",
        "visual_headers": ["Visual", "Type", "Fields / references"],
        "no_visuals": "No visuals found.",
        "lineage_banner": "Complete dependency graph for the pipeline",
        "lineage_headers": ["Source", "Relationship", "Target"],
        "no_lineage": "No lineage edges found.",
        "diagnostics_banner": "Warnings and technical information from the analysis",
        "relationship_columns": "Available columns in the relationships table:",
        "rel_from": "From",
        "rel_to": "To",
        "cardinality": "Cardinality",
        "cross_filter": "Cross filter",
        "relationship_example": "Detected relationship example — ",
        "message": "Message",
        "no_diagnostics": "No diagnostics recorded.",
        "yes": "Yes",
        "no": "No",
        "dash": "—",
        "m_code_heading": "3.1 M code by query",
        "summary": "Summary",
        "technical_expressions": "Technical expressions",
        "additional_records_omitted": "{count} additional records were omitted in this section to keep the document readable.",
        "no_records": "No records found.",
        "record": "Record",
        "truncated": "truncated",
        "connector": "Connector",
        "steps": "Steps",
        "functions": "Functions",
        "datatype_map": {
            "2": "text", "3": "decimal", "4": "integer", "5": "decimal",
            "6": "currency", "7": "date", "8": "boolean", "9": "binary",
            "10": "variant", "17": "integer",
            "string": "text", "int64": "integer", "double": "decimal",
            "boolean": "boolean", "datetime": "date/time", "binary": "binary",
        },
    },
}


def normalize_doc_locale(locale: str = "") -> str:
    normalized = str(locale or "").strip().lower()
    if normalized.startswith("en"):
        return "en-US"
    return "pt-BR"


def doc_text(locale: str = ""):
    return DOC_TEXT[normalize_doc_locale(locale)]


def build_source_rows(sources, labels):
    """Linhas [Fonte, Padrao detectado, Connection string] da secao "Fontes
    de dados"."""
    rows = []
    for source in sources:
        meta = source.get("meta", {})
        datasource_meta = meta.get("datasource") or {}
        connection_string = ""
        if isinstance(datasource_meta, dict):
            connection_string = (
                datasource_meta.get("ConnectionString", "")
                or datasource_meta.get("Location", "")
                or datasource_meta.get("Account", "")
                or datasource_meta.get("Path", "")
                or ""
            )
        if not connection_string:
            connection_string = meta.get("connectionPath") or meta.get("doc") or ""
        rows.append([
            source.get("label", ""),
            meta.get("pattern", ""),
            mask_secrets(connection_string) if connection_string else labels["dash"],
        ])
    return rows


def build_datasource_rows(datasource_records, labels, limit=30):
    """Linhas [Tipo, Nome, Connection string] de TMSCHEMA_DATASOURCES cru
    (subsecao "2.1 Detalhes de conexao")."""
    rows = []
    for ds in datasource_records[:limit]:
        conn_str = (
            ds.get("ConnectionString", "")
            or ds.get("Location", "")
            or ds.get("Account", "")
            or ds.get("Path", "")
            or ""
        )
        kind = ds.get("Kind", ds.get("Type", ds.get("SourceType", "")))
        name = ds.get("Name", ds.get("ContentPath", ""))
        rows.append([
            str(kind) if kind else labels["dash"],
            str(name) if name else labels["dash"],
            mask_secrets(conn_str) if conn_str else labels["dash"],
        ])
    return rows


def build_query_rows(queries, labels):
    """Linhas [Query, Conexao/caminho, Resumo tecnico] da secao "Power
    Query"."""
    rows = []
    for query in queries:
        meta = query.get("meta", {})
        rows.append([
            query.get("label", ""),
            mask_secrets(meta.get("connectionPath", "")),
            summarize_m_expression(mask_secrets(meta.get("expression", "")), labels),
        ])
    return rows


def build_relationship_rows(relationships, labels):
    """Linhas da secao "Relacionamentos"."""
    return [[
        rel.get("fromTable", ""),
        rel.get("fromColumn", ""),
        rel.get("toTable", ""),
        rel.get("toColumn", ""),
        rel.get("cardinality", ""),
        rel.get("crossFilter", ""),
        labels["yes"] if rel.get("active", True) else labels["no"],
    ] for rel in relationships]


def build_page_rows(pages):
    """Linhas [#, Pagina, Visuais, Canvas] da secao "Paginas e visuais"."""
    return [
        [page.get("ordinal", 0) + 1, page.get("name", ""), page.get("visualCount", 0),
         f"{page.get('width', '')} x {page.get('height', '')}"]
        for page in pages
    ]


def build_visual_rows(visuals, labels):
    """Linhas [Visual, Tipo, Campos/referencias], deduplicando refs ruidosas."""
    rows = []
    for visual in visuals:
        meta = visual.get("meta", {})
        refs_raw = meta.get("refs", []) or []
        seen_refs = set()
        clean_refs = []
        for ref in refs_raw:
            ref_str = str(ref).strip()
            if len(ref_str) < 2:
                continue
            ref_lower = ref_str.lower()
            if ref_lower in seen_refs:
                continue
            seen_refs.add(ref_lower)
            clean_refs.append(ref_str)
        refs_display = "\n".join(clean_refs) if clean_refs else labels["dash"]
        rows.append([visual.get("label", ""), meta.get("visualType", ""), refs_display])
    return rows


def build_edge_rows(nodes, edges):
    """Linhas [Origem, Relacao, Destino] da secao "Linhagem tecnica"."""
    node_labels = {node.get("id"): node.get("label", node.get("id", "")) for node in nodes}
    return [
        [node_labels.get(item.get("from"), item.get("from", "")), item.get("label", ""),
         node_labels.get(item.get("to"), item.get("to", ""))]
        for item in edges
    ]


def build_readable_warnings(warnings, labels):
    """Reformata as strings de diagnostico internas (ex. "relationships
    columns: [...]") em frases legiveis para a secao "Diagnosticos"."""
    readable = []
    for w in unique_texts(warnings):
        if not w:
            continue
        if w.startswith("relationships columns:"):
            cols_match = re.search(r"\[(.+)\]", w)
            if cols_match:
                cols = [c.strip().strip("'") for c in cols_match.group(1).split(",")]
                readable.append(f"{labels['relationship_columns']} {', '.join(cols)}")
            else:
                readable.append(w)
        elif w.startswith("relationships first row:"):
            try:
                dict_str = w[len("relationships first row: "):]
                row_dict = {}
                for m in re.finditer(r"'([^']+)':\s*(?:'([^']*)'|(\d+)|None)", dict_str):
                    key, sv, nv = m.group(1), m.group(2), m.group(3)
                    row_dict[key] = sv if sv is not None else (nv if nv is not None else "")
                parts = []
                if row_dict.get("FromTableName"):
                    parts.append(f"{labels['rel_from']}: {row_dict['FromTableName']}[{row_dict.get('FromColumnName', '')}]")
                if row_dict.get("ToTableName"):
                    parts.append(f"{labels['rel_to']}: {row_dict['ToTableName']}[{row_dict.get('ToColumnName', '')}]")
                if row_dict.get("Cardinality"):
                    parts.append(f"{labels['cardinality']}: {row_dict['Cardinality']}")
                if row_dict.get("CrossFilteringBehavior"):
                    parts.append(f"{labels['cross_filter']}: {row_dict['CrossFilteringBehavior']}")
                if parts:
                    readable.append(labels["relationship_example"] + " | ".join(parts))
                else:
                    readable.append(w)
            except Exception:
                logger.debug("Falha ao formatar diagnostico legivel: %r", w, exc_info=True)
                readable.append(w)
        else:
            readable.append(w)
    return readable


def build_documentation_docx(path: Path, file_name: str = "", locale: str = "pt-BR") -> bytes:
    import backend  # adiado -- ver docstring do modulo / pbix_analysis.py

    if not backend.load_pbixray():
        raise RuntimeError(f"pbixray nao esta instalado ou nao carregou: {backend.PBIXRAY_IMPORT_ERROR}")
    if not load_docx():
        raise RuntimeError(f"python-docx nao esta instalado ou nao carregou: {DOCX_IMPORT_ERROR}")
    load_pillow()

    doc_labels = doc_text(locale)
    metric_labels = doc_labels["metrics"]
    diagnostics = []
    model = backend.PBIXRay(str(path))
    graph = analyze_pbix(path)

    power_query = records_from(model, "power_query", diagnostics)
    datasource_records = records_from(model, "tmschema_datasources", diagnostics)
    measures = records_from(model, "dax_measures", diagnostics)
    calc_columns = records_from(model, "dax_columns", diagnostics)
    schema = records_from(model, "schema", diagnostics)
    semantic_tables = records_from(model, "tmschema_tables", diagnostics)
    tmschema_columns = records_from(model, "tmschema_columns", diagnostics)

    nodes = graph.get("nodes", [])
    edges = graph.get("edges", [])
    relationships = graph.get("relationships", [])
    pages = graph.get("pages", [])
    warnings = list(graph.get("warnings", [])) + diagnostics

    sources = [node for node in nodes if node.get("type") == "source"]
    queries = [node for node in nodes if node.get("type") == "query"]
    tables = [node for node in nodes if node.get("type") == "model"]
    measure_nodes = [node for node in nodes if node.get("type") == "measure"]
    calc_column_nodes = [node for node in nodes if node.get("type") == "calc_column"]
    visuals = [node for node in nodes if node.get("type") == "visual"]

    doc = Document()
    configure_document_styles(doc)
    add_cover(doc, file_name or path.name, {
        metric_labels["sources"]: len(sources),
        metric_labels["queries"]: len(queries),
        metric_labels["tables"]: len(tables),
        metric_labels["measures"]: len(measure_nodes),
        metric_labels["calc_columns"]: len(calc_column_nodes),
        metric_labels["relationships"]: len(relationships),
        metric_labels["pages"]: len(pages),
        metric_labels["visuals"]: len(visuals),
    }, doc_labels)
    add_table_of_contents(doc, doc_labels["sections"], doc_labels)

    add_doc_heading(doc, f"1. {doc_labels['sections'][0]}", level=1)
    add_section_banner(doc, doc_labels["overview_banner"], color="#0078D4", icon="📋")
    add_paragraph(
        doc,
        doc_labels["overview_text"],
    )
    add_key_value_table(doc, [
        (doc_labels["file"], file_name or path.name),
        (doc_labels["generated_at"], datetime.now().strftime("%d/%m/%Y %H:%M:%S")),
        (doc_labels["analysis_source"], "PBIXRay"),
        (doc_labels["total_nodes"], len(nodes)),
        (doc_labels["total_edges"], len(edges)),
    ])

    add_doc_heading(doc, f"2. {doc_labels['sections'][1]}", level=1)
    add_section_banner(doc, doc_labels["data_sources_banner"], color="#0078D4", icon="🔌")
    source_rows = build_source_rows(sources, doc_labels)
    add_records_table(doc, doc_labels["source_headers"], source_rows, empty=doc_labels["no_sources"], labels=doc_labels)

    # ── Diagrama de arquitetura ──────────────────────────────────────────────
    arch_png = make_arch_png(sources, queries, edges)
    if arch_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(BytesIO(arch_png), width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(doc_labels["arch_caption"])
        cr.font.size = Pt(8.5); cr.font.italic = True; cr.font.color.rgb = RGBColor(96,94,92)
    elif sources:
        add_paragraph(doc, f"{len(sources)} {doc_labels['sources_detected']} " + ", ".join(s.get('label','') for s in sources))

    if datasource_records:
        add_doc_heading(doc, doc_labels["connection_details"], level=2)
        arch_rows = build_datasource_rows(datasource_records, doc_labels, limit=30)
        if arch_rows:
            add_records_table(doc, doc_labels["connection_headers"], arch_rows, empty="", labels=doc_labels)
        else:
            # G7: datasource_records cru pode conter ConnectionString/Pwd
            # verbatim em qualquer uma das suas colunas -- mascara tudo, nao
            # so os campos que a extracao estruturada acima ja reconhece.
            add_generic_records_table(doc, datasource_records, max_rows=30, labels=doc_labels, mask=True)

    add_doc_heading(doc, f"3. {doc_labels['sections'][2]}", level=1)
    add_section_banner(doc, doc_labels["power_query_banner"], color="#F2C811", icon="⚙️")
    query_rows = build_query_rows(queries, doc_labels)
    add_records_table(doc, doc_labels["query_headers"], query_rows, empty=doc_labels["no_queries"], labels=doc_labels)

    add_power_query_evidence(doc, power_query, doc_labels)

    add_doc_heading(doc, f"4. {doc_labels['sections'][3]}", level=1)
    add_section_banner(doc, doc_labels["semantic_model_banner"], color="#107C10", icon="🗄️")
    table_rows = build_table_documentation_rows(tables, semantic_tables, schema, doc_labels)
    add_records_table(doc, doc_labels["table_headers"], table_rows, empty=doc_labels["no_tables"], labels=doc_labels)

    add_doc_heading(doc, f"5. {doc_labels['sections'][4]}", level=1)
    add_section_banner(doc, doc_labels["dictionary_banner"], color="#107C10", icon="📖")
    column_rows = build_column_rows(schema, tmschema_columns, semantic_tables, doc_labels)
    add_records_table(
        doc,
        doc_labels["column_headers"],
        column_rows,
        empty=doc_labels["no_columns"],
        max_rows=250,
        labels=doc_labels,
    )

    add_doc_heading(doc, f"6. {doc_labels['sections'][5]}", level=1)
    add_section_banner(doc, doc_labels["measures_banner"], color="#D83B01", icon="📐")
    measure_rows = build_expression_rows(measures, measure_nodes, kind="measure")
    add_expression_inventory(doc, doc_labels["measure_headers"], measure_rows, empty=doc_labels["no_measures"], labels=doc_labels)

    add_doc_heading(doc, f"7. {doc_labels['sections'][6]}", level=1)
    add_section_banner(doc, doc_labels["calc_columns_banner"], color="#9B5094", icon="🔢")
    calc_rows = build_expression_rows(calc_columns, calc_column_nodes, kind="calc_column")
    add_expression_inventory(doc, doc_labels["calc_headers"], calc_rows, empty=doc_labels["no_calc_columns"], labels=doc_labels)

    add_doc_heading(doc, f"8. {doc_labels['sections'][7]}", level=1)
    add_section_banner(doc, doc_labels["relationships_banner"], color="#0078D4", icon="🔗")
    relationship_rows = build_relationship_rows(relationships, doc_labels)
    add_records_table(
        doc,
        doc_labels["relationship_headers"],
        relationship_rows,
        empty=doc_labels["no_relationships"],
        max_rows=200,
        labels=doc_labels,
    )

    # ── Diagrama ERD ─────────────────────────────────────────────────────────
    erd_png = make_erd_png(relationships)
    if erd_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(BytesIO(erd_png), width=Inches(6.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(doc_labels["erd_caption"])
        cr.font.size = Pt(8.5); cr.font.italic = True; cr.font.color.rgb = RGBColor(96,94,92)

    add_doc_heading(doc, f"9. {doc_labels['sections'][8]}", level=1)
    add_section_banner(doc, doc_labels["pages_banner"], color="#8764B8", icon="📄")
    page_rows = build_page_rows(pages)
    add_records_table(doc, doc_labels["page_headers"], page_rows, empty=doc_labels["no_pages"], labels=doc_labels)

    visual_rows = build_visual_rows(visuals, doc_labels)
    add_records_table(doc, doc_labels["visual_headers"], visual_rows, empty=doc_labels["no_visuals"], max_rows=120, labels=doc_labels)

    add_doc_heading(doc, f"10. {doc_labels['sections'][9]}", level=1)
    add_section_banner(doc, doc_labels["lineage_banner"], color="#1B2A38", icon="🔀")
    edge_rows = build_edge_rows(nodes, edges)
    add_records_table(doc, doc_labels["lineage_headers"], edge_rows, empty=doc_labels["no_lineage"], max_rows=300, labels=doc_labels)

    add_doc_heading(doc, f"11. {doc_labels['sections'][10]}", level=1)
    add_section_banner(doc, doc_labels["diagnostics_banner"], color="#605E5C", icon="🔍")
    readable_warnings = build_readable_warnings(warnings, doc_labels)
    warning_rows = [[w] for w in readable_warnings if w]
    add_records_table(doc, [doc_labels["message"]], warning_rows, empty=doc_labels["no_diagnostics"], max_rows=120, labels=doc_labels)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _html_escape(value) -> str:
    return html.escape(doc_value(value), quote=True)


def _html_document_css() -> str:
    return """
    :root { color-scheme: light; }
    * { box-sizing: border-box; }
    body {
      font-family: "Segoe UI", Aptos, Arial, sans-serif;
      color: #1F2937; background: #F3F2F1; margin: 0;
      line-height: 1.45;
    }
    .page { max-width: 980px; margin: 0 auto; background: #FFFFFF; padding: 32px 40px 64px; }
    .cover { text-align: center; padding: 28px 0 20px; border-bottom: 4px solid #F2C811; margin-bottom: 24px; }
    .cover h1 { color: #0B1F33; font-size: 30px; margin: 8px 0 2px; }
    .cover .subtitle { color: #1E6D85; font-size: 14px; margin: 0 0 18px; }
    .cover .filename { font-weight: 700; font-size: 15px; color: #233040; margin-bottom: 18px; }
    .metric-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px; margin: 18px 0; }
    .metric { background: #F4F7FA; border-radius: 6px; padding: 14px 8px; text-align: center; }
    .metric-value { font-weight: 700; font-size: 20px; color: #0B1F33; }
    .metric-label { font-size: 10px; color: #586776; margin-top: 2px; }
    h2.section-title {
      color: #0B1F33; font-size: 19px; border-bottom: 2px solid #D8DEE8;
      padding-bottom: 6px; margin: 34px 0 10px;
    }
    h3 { color: #1E6D85; font-size: 14px; margin: 18px 0 6px; }
    .banner {
      font-weight: 600; font-size: 13px; padding: 10px 14px; border-radius: 4px;
      margin: 6px 0 14px;
    }
    table { border-collapse: collapse; width: 100%; margin: 10px 0 18px; font-size: 12px; }
    table.kv-table { width: auto; min-width: 45%; }
    th, td { border: 1px solid #D8DEE8; padding: 6px 9px; text-align: left; vertical-align: top; }
    thead th { background: #0B1F33; color: #FFFFFF; font-weight: 600; }
    table.kv-table th { background: #F4F7FA; color: #1F2937; width: 220px; }
    tbody tr:nth-child(even) { background: #F8FAFC; }
    .muted { color: #6B7280; font-style: italic; font-size: 12.5px; }
    .code-block { margin: 10px 0 16px; }
    .code-title { font-weight: 700; color: #1E6D85; font-size: 12.5px; margin-bottom: 3px; }
    .code-block pre {
      background: #F6F8FA; border: 1px solid #D8DEE8; border-radius: 4px;
      padding: 10px 12px; font-family: Consolas, "Courier New", monospace;
      font-size: 11px; white-space: pre-wrap; word-break: break-word; margin: 0;
    }
    ol.toc { columns: 1; padding-left: 22px; }
    ol.toc li { margin: 4px 0; }
    ol.toc a { color: #1E6D85; text-decoration: none; }
    ol.toc a:hover { text-decoration: underline; }
    figure.diagram { margin: 14px 0 20px; text-align: center; }
    figure.diagram img { max-width: 100%; height: auto; border: 1px solid #E1DFDD; border-radius: 4px; }
    figure.diagram figcaption { font-size: 10.5px; font-style: italic; color: #605E5C; margin-top: 6px; }
    footer.doc-footer { text-align: right; color: #778696; font-size: 10.5px; margin-top: 30px; }
    """


def _html_metric_grid(metrics: dict) -> str:
    cells = "".join(
        f'<div class="metric"><div class="metric-value">{_html_escape(value)}</div>'
        f'<div class="metric-label">{_html_escape(label)}</div></div>'
        for label, value in metrics.items()
    )
    return f'<div class="metric-grid">{cells}</div>'


def _html_key_value_table(rows) -> str:
    body = "".join(f"<tr><th>{_html_escape(k)}</th><td>{_html_escape(v)}</td></tr>" for k, v in rows)
    return f'<table class="kv-table"><tbody>{body}</tbody></table>'


def _html_toc(sections, labels) -> str:
    items = "".join(
        f'<li><a href="#section-{index}">{index:02d}. {_html_escape(title)}</a></li>'
        for index, title in enumerate(sections, 1)
    )
    return f'<h2 class="section-title">{_html_escape(labels["toc"])}</h2><ol class="toc">{items}</ol>'


def _html_banner(text: str, color: str) -> str:
    return (
        f'<div class="banner" style="background:{color}1a;border-left:4px solid {color};color:{color};">'
        f'{_html_escape(text)}</div>'
    )


def _html_table(headers, rows, empty_text, labels, max_rows=200) -> str:
    if not rows:
        return f'<p class="muted">{_html_escape(empty_text)}</p>'
    limited_rows = rows[:max_rows]
    thead = "".join(f"<th>{_html_escape(h)}</th>" for h in headers)
    body_rows = []
    for row in limited_rows:
        cells = "".join(
            f"<td>{_html_escape(value).replace(chr(10), '<br>')}</td>"
            for value in row[: len(headers)]
        )
        body_rows.append(f"<tr>{cells}</tr>")
    extra = ""
    if len(rows) > max_rows:
        extra = f'<p class="muted">{_html_escape(labels["additional_records_omitted"].format(count=len(rows) - max_rows))}</p>'
    return f'<table><thead><tr>{thead}</tr></thead><tbody>{"".join(body_rows)}</tbody></table>{extra}'


def _html_code_block(title: str, code: str, labels, max_chars=6000) -> str:
    body = normalize_code_text(clip(code, max_chars, labels))
    return (
        f'<div class="code-block"><div class="code-title">{_html_escape(title)}</div>'
        f'<pre>{_html_escape(body)}</pre></div>'
    )


def _html_expression_inventory(headers, rows, empty_text, labels) -> str:
    inventory_rows = [[row[0], row[1], summarize_expression(row[2], labels)] for row in rows]
    parts = [_html_table(headers + [labels["summary"]], inventory_rows, empty_text, labels, max_rows=200)]
    expressive_rows = [row for row in rows if str(row[2] or "").strip()]
    if expressive_rows:
        parts.append(f'<h3>{_html_escape(labels["technical_expressions"])}</h3>')
        for table, name, expression in expressive_rows[:80]:
            title = f"{table} - {name}" if table else str(name)
            parts.append(_html_code_block(title, expression, labels, max_chars=2200))
    return "".join(parts)


def _html_svg_figure(svg_markup: str, caption: str) -> str:
    if not svg_markup:
        return ""
    # Base64 data-URI dentro de <img>: o navegador renderiza o SVG como
    # imagem opaca (nao executa <script>/handlers embutidos, ao contrario de
    # SVG inline no DOM) -- ver nota de seguranca no topo desta secao.
    encoded = base64.b64encode(svg_markup.encode("utf-8")).decode("ascii")
    return (
        '<figure class="diagram">'
        f'<img alt="{_html_escape(caption)}" src="data:image/svg+xml;base64,{encoded}">'
        f'<figcaption>{_html_escape(caption)}</figcaption>'
        '</figure>'
    )


def build_documentation_html(path: Path, file_name: str = "", locale: str = "pt-BR") -> bytes:
    import backend  # adiado -- ver docstring do modulo / pbix_analysis.py

    if not backend.load_pbixray():
        raise RuntimeError(f"pbixray nao esta instalado ou nao carregou: {backend.PBIXRAY_IMPORT_ERROR}")

    doc_labels = doc_text(locale)
    metric_labels = doc_labels["metrics"]
    diagnostics = []
    model = backend.PBIXRay(str(path))
    graph = analyze_pbix(path)

    power_query = records_from(model, "power_query", diagnostics)
    datasource_records = records_from(model, "tmschema_datasources", diagnostics)
    measures = records_from(model, "dax_measures", diagnostics)
    calc_columns = records_from(model, "dax_columns", diagnostics)
    schema = records_from(model, "schema", diagnostics)
    semantic_tables = records_from(model, "tmschema_tables", diagnostics)
    tmschema_columns = records_from(model, "tmschema_columns", diagnostics)

    nodes = graph.get("nodes", [])
    edges = graph.get("edges", [])
    relationships = graph.get("relationships", [])
    pages = graph.get("pages", [])
    warnings = list(graph.get("warnings", [])) + diagnostics

    sources = [node for node in nodes if node.get("type") == "source"]
    queries = [node for node in nodes if node.get("type") == "query"]
    tables = [node for node in nodes if node.get("type") == "model"]
    measure_nodes = [node for node in nodes if node.get("type") == "measure"]
    calc_column_nodes = [node for node in nodes if node.get("type") == "calc_column"]
    visuals = [node for node in nodes if node.get("type") == "visual"]

    display_name = file_name or path.name
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    metrics = {
        metric_labels["sources"]: len(sources),
        metric_labels["queries"]: len(queries),
        metric_labels["tables"]: len(tables),
        metric_labels["measures"]: len(measure_nodes),
        metric_labels["calc_columns"]: len(calc_column_nodes),
        metric_labels["relationships"]: len(relationships),
        metric_labels["pages"]: len(pages),
        metric_labels["visuals"]: len(visuals),
    }

    parts = [
        "<!DOCTYPE html>",
        f'<html lang="{"pt-BR" if locale == "pt-BR" else "en-US"}">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>BI Flow Mapper — {_html_escape(display_name)}</title>",
        f"<style>{_html_document_css()}</style>",
        "</head>",
        "<body><div class='page'>",
        '<div class="cover">',
        '<h1>BI Flow Mapper</h1>',
        f'<div class="subtitle">{_html_escape(doc_labels["cover_subtitle"])}</div>',
        f'<div class="filename">{_html_escape(display_name)}</div>',
        _html_metric_grid(metrics),
        f'<div class="muted">{_html_escape(doc_labels["generated_at"])} {generated_at}</div>',
        "</div>",
        _html_toc(doc_labels["sections"], doc_labels),
    ]

    def section(number, title, banner_text, color, body_html):
        parts.append(f'<h2 class="section-title" id="section-{number}">{number}. {_html_escape(title)}</h2>')
        parts.append(_html_banner(banner_text, color))
        parts.append(body_html)

    # 1. Visão geral
    overview_body = (
        f'<p>{_html_escape(doc_labels["overview_text"])}</p>'
        + _html_key_value_table([
            (doc_labels["file"], display_name),
            (doc_labels["generated_at"], generated_at),
            (doc_labels["analysis_source"], "PBIXRay"),
            (doc_labels["total_nodes"], len(nodes)),
            (doc_labels["total_edges"], len(edges)),
        ])
    )
    section(1, doc_labels["sections"][0], doc_labels["overview_banner"], "#0078D4", overview_body)

    # 2. Fontes de dados
    source_rows = build_source_rows(sources, doc_labels)
    sources_body = _html_table(doc_labels["source_headers"], source_rows, doc_labels["no_sources"], doc_labels)
    arch_svg = build_architecture_svg(sources, queries, edges)
    sources_body += _html_svg_figure(arch_svg, doc_labels["arch_caption"])
    if datasource_records:
        sources_body += f'<h3>{_html_escape(doc_labels["connection_details"])}</h3>'
        datasource_rows = build_datasource_rows(datasource_records, doc_labels, limit=30)
        sources_body += _html_table(doc_labels["connection_headers"], datasource_rows, "", doc_labels)
    section(2, doc_labels["sections"][1], doc_labels["data_sources_banner"], "#0078D4", sources_body)

    # 3. Power Query
    query_rows = build_query_rows(queries, doc_labels)
    pq_body = _html_table(doc_labels["query_headers"], query_rows, doc_labels["no_queries"], doc_labels)
    if power_query:
        pq_body += f'<h3>{_html_escape(doc_labels["m_code_heading"])}</h3>'
        for index, row in enumerate(power_query[:30], 1):
            name = first_value(row, ["TableName", "Name", "QueryName", "table", "name"]) or f"Query {index}"
            expression = first_value(row, ["Expression", "expression", "MExpression"]) or record_text(row)
            pq_body += _html_code_block(str(name), mask_secrets(expression), doc_labels, max_chars=2800)
    section(3, doc_labels["sections"][2], doc_labels["power_query_banner"], "#F2C811", pq_body)

    # 4. Modelo semântico
    table_rows = build_table_documentation_rows(tables, semantic_tables, schema, doc_labels)
    section(4, doc_labels["sections"][3], doc_labels["semantic_model_banner"], "#107C10",
            _html_table(doc_labels["table_headers"], table_rows, doc_labels["no_tables"], doc_labels))

    # 5. Dicionário de dados
    column_rows = build_column_rows(schema, tmschema_columns, semantic_tables, doc_labels)
    section(5, doc_labels["sections"][4], doc_labels["dictionary_banner"], "#107C10",
            _html_table(doc_labels["column_headers"], column_rows, doc_labels["no_columns"], doc_labels, max_rows=250))

    # 6. Medidas DAX
    measure_rows = build_expression_rows(measures, measure_nodes, kind="measure")
    section(6, doc_labels["sections"][5], doc_labels["measures_banner"], "#D83B01",
            _html_expression_inventory(doc_labels["measure_headers"], measure_rows, doc_labels["no_measures"], doc_labels))

    # 7. Colunas calculadas
    calc_rows = build_expression_rows(calc_columns, calc_column_nodes, kind="calc_column")
    section(7, doc_labels["sections"][6], doc_labels["calc_columns_banner"], "#9B5094",
            _html_expression_inventory(doc_labels["calc_headers"], calc_rows, doc_labels["no_calc_columns"], doc_labels))

    # 8. Relacionamentos
    relationship_rows = build_relationship_rows(relationships, doc_labels)
    rel_body = _html_table(doc_labels["relationship_headers"], relationship_rows, doc_labels["no_relationships"], doc_labels, max_rows=200)
    erd_svg = build_erd_svg(relationships)
    rel_body += _html_svg_figure(erd_svg, doc_labels["erd_caption"])
    section(8, doc_labels["sections"][7], doc_labels["relationships_banner"], "#0078D4", rel_body)

    # 9. Páginas e visuais
    page_rows = build_page_rows(pages)
    pages_body = _html_table(doc_labels["page_headers"], page_rows, doc_labels["no_pages"], doc_labels)
    visual_rows = build_visual_rows(visuals, doc_labels)
    pages_body += _html_table(doc_labels["visual_headers"], visual_rows, doc_labels["no_visuals"], doc_labels, max_rows=120)
    section(9, doc_labels["sections"][8], doc_labels["pages_banner"], "#8764B8", pages_body)

    # 10. Linhagem técnica
    edge_rows = build_edge_rows(nodes, edges)
    section(10, doc_labels["sections"][9], doc_labels["lineage_banner"], "#1B2A38",
            _html_table(doc_labels["lineage_headers"], edge_rows, doc_labels["no_lineage"], doc_labels, max_rows=300))

    # 11. Diagnósticos
    readable_warnings = build_readable_warnings(warnings, doc_labels)
    warning_rows = [[w] for w in readable_warnings if w]
    section(11, doc_labels["sections"][10], doc_labels["diagnostics_banner"], "#605E5C",
            _html_table([doc_labels["message"]], warning_rows, doc_labels["no_diagnostics"], doc_labels, max_rows=120))

    parts.append(f'<footer class="doc-footer">BI Flow Mapper — {_html_escape(generated_at)}</footer>')
    parts.append("</div></body></html>")

    return "".join(parts).encode("utf-8")


def configure_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("BI Flow Mapper")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(119, 134, 150)

    for style_name, size, color in [
        ("Title", 24, "0B1F33"),
        ("Heading 1", 15, "0B1F33"),
        ("Heading 2", 11, "1E6D85"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_cover(doc, file_name, metrics, labels=None):
    labels = labels or doc_text("pt-BR")
    # ── Gold top stripe (Power BI identity) ─────────────────────────────────
    def _make_stripe() -> bytes | None:
        if not render_graphics.PIL_AVAILABLE:
            return None
        try:
            img = render_graphics._PILImage.new("RGB", (1440, 36), (27, 42, 56))
            draw = render_graphics._PILDraw.Draw(img)
            draw.rectangle([0, 24, 1440, 36], fill=(242, 200, 17))
            return _pil_png(img)
        except Exception:
            logger.debug("_make_stripe falhou -- capa do DOCX sera exibida sem a faixa dourada.", exc_info=True)
            return None
    stripe_png = _make_stripe()
    if stripe_png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(BytesIO(stripe_png), width=Inches(7.0))

    logo = ROOT / "image" / "icon.png"
    if logo.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.paragraph_format.space_before = Pt(6)
        logo_paragraph.paragraph_format.space_after = Pt(0)
        logo_paragraph.add_run().add_picture(str(logo), width=Inches(0.65))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("BI Flow Mapper")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(11, 31, 51)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    subtitle_run = subtitle.add_run(labels["cover_subtitle"])
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(30, 109, 133)

    file_paragraph = doc.add_paragraph()
    file_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file_paragraph.paragraph_format.space_after = Pt(22)
    file_run = file_paragraph.add_run(file_name)
    file_run.bold = True
    file_run.font.size = Pt(15)
    file_run.font.color.rgb = RGBColor(35, 48, 64)

    add_metric_grid(doc, metrics)
    add_paragraph(doc, f"{labels['generated_at']} {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    # ── Gold bottom stripe ───────────────────────────────────────────────────
    if stripe_png:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(10)
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run().add_picture(BytesIO(stripe_png), width=Inches(7.0))

    doc.add_page_break()


def add_table_of_contents(doc, sections, labels=None):
    labels = labels or doc_text("pt-BR")
    add_doc_heading(doc, labels["toc"], level=1)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    set_table_width(table, [0.55, 6.2])
    for index, title in enumerate(sections, 1):
        cells = table.add_row().cells
        cells[0].text = f"{index:02d}"
        cells[1].text = title
        for cell in cells:
            set_cell_margins(cell, top=80, start=80, bottom=80, end=80)
            if index % 2 == 0:
                shade_cell(cell, "F8FAFC")
        set_cell_font(cells[0], color="1E6D85", bold=True)
        set_cell_font(cells[1], color="1F2937", bold=True)
    doc.add_page_break()


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(str(text or ""))
    return paragraph


def add_section_banner(doc, text: str, color: str = "#0078D4", icon: str = ""):
    """Insert a colored banner below a heading for visual identity."""
    png = make_banner_png(text, color=color, icon=icon)
    if not png:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(BytesIO(png), width=Inches(6.8))


def add_doc_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    run = paragraph.add_run(str(text))
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(15 if level == 1 else 11)
    run.font.color.rgb = RGBColor(11, 31, 51) if level == 1 else RGBColor(30, 109, 133)
    if level == 1:
        add_paragraph_bottom_border(paragraph, "D8DEE8", "8")
    return paragraph


def add_metric_grid(doc, metrics):
    items = list(metrics.items())
    cols = 4
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF")
    for row_index in range(rows):
        for col_index in range(cols):
            cell = table.rows[row_index].cells[col_index]
            set_cell_margins(cell, top=140, start=140, bottom=140, end=140)
            shade_cell(cell, "F4F7FA")
            item_index = row_index * cols + col_index
            if item_index >= len(items):
                cell.text = ""
                continue
            label, value = items[item_index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            value_run = paragraph.add_run(str(value))
            value_run.bold = True
            value_run.font.size = Pt(18)
            value_run.font.color.rgb = RGBColor(11, 31, 51)
            paragraph.add_run("\n")
            label_run = paragraph.add_run(str(label))
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = RGBColor(88, 103, 118)
    doc.add_paragraph("")


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_borders(table, color="D8DEE8")
    set_table_width(table, [2.1, 4.9])
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
        shade_cell(cells[0], "F4F7FA")
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        bold_cell(cells[0])
    doc.add_paragraph("")


def add_records_table(doc, headers, rows, empty="", max_rows=120, labels=None):
    labels = labels or doc_text("pt-BR")
    if not rows:
        add_paragraph(doc, empty)
        return

    limited_rows = rows[:max_rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    set_table_borders(table, color="D8DEE8")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        shade_cell(cell, "0B1F33")
        set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
        set_cell_font(cell, color="FFFFFF", bold=True)

    for row_index, row in enumerate(limited_rows):
        cells = table.add_row().cells
        for index, value in enumerate(row[:len(headers)]):
            cells[index].text = clip(value, 700, labels)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[index], top=70, start=90, bottom=70, end=90)
            if row_index % 2:
                shade_cell(cells[index], "F8FAFC")
            set_cell_font(cells[index], color="1F2937")

    if len(rows) > max_rows:
        add_paragraph(doc, labels["additional_records_omitted"].format(count=len(rows) - max_rows))
    doc.add_paragraph("")


def add_generic_records_table(doc, records, max_rows=60, labels=None, mask=False):
    labels = labels or doc_text("pt-BR")
    keys = []
    for record in records:
        for key in record.keys():
            if key not in keys:
                keys.append(key)
        if len(keys) >= 6:
            break
    value_fn = mask_secrets if mask else doc_value
    rows = [[value_fn(record.get(key, "")) for key in keys] for record in records]
    add_records_table(doc, keys or [labels["record"]], rows, empty=labels["no_records"], max_rows=max_rows, labels=labels)


def add_power_query_evidence(doc, power_query, labels=None):
    labels = labels or doc_text("pt-BR")
    if not power_query:
        return
    add_doc_heading(doc, labels["m_code_heading"], level=2)
    for index, row in enumerate(power_query[:30], 1):
        name = first_value(row, ["TableName", "Name", "QueryName", "table", "name"]) or f"Query {index}"
        expression = first_value(row, ["Expression", "expression", "MExpression"]) or record_text(row)
        add_code_block(doc, str(name), mask_secrets(expression), max_chars=2800, labels=labels)


def add_expression_inventory(doc, headers, rows, empty="", labels=None):
    labels = labels or doc_text("pt-BR")
    if not rows:
        add_paragraph(doc, empty)
        return

    inventory_rows = [[row[0], row[1], summarize_expression(row[2], labels)] for row in rows]
    add_records_table(doc, headers + [labels["summary"]], inventory_rows, empty=empty, max_rows=200, labels=labels)

    expressive_rows = [row for row in rows if str(row[2] or "").strip()]
    if not expressive_rows:
        return

    add_doc_heading(doc, labels["technical_expressions"], level=2)
    for table, name, expression in expressive_rows[:80]:
        title = f"{name}"
        if table:
            title = f"{table} - {name}"
        add_code_block(doc, title, expression, max_chars=2200, labels=labels)


def add_code_block(doc, title, code, max_chars=2600, labels=None):
    labels = labels or doc_text("pt-BR")
    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.keep_with_next = True
    title_paragraph.paragraph_format.space_before = Pt(8)
    title_paragraph.paragraph_format.space_after = Pt(3)
    run = title_paragraph.add_run(str(title))
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 109, 133)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="D8DEE8")
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F6F8FA")
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    code_run = paragraph.add_run(normalize_code_text(clip(code, max_chars, labels)))
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(7.5)
    code_run.font.color.rgb = RGBColor(31, 41, 55)
    doc.add_paragraph("")


def build_table_documentation_rows(table_nodes, semantic_tables, schema, labels=None):
    labels = labels or doc_text("pt-BR")
    table_meta = {}
    for row in semantic_tables:
        name = first_value(row, ["Name", "TableName", "Table", "name"])
        if name:
            table_meta[str(name)] = row

    column_counts = {}
    for row in schema:
        table = first_value(row, ["TableName", "Table", "table"])
        column = first_value(row, ["ColumnName", "Name", "Column", "column", "name"])
        if table and column:
            column_counts[str(table)] = column_counts.get(str(table), 0) + 1

    rows = []
    for node in table_nodes:
        name = node.get("label", "")
        meta = table_meta.get(name, {})
        rows.append([
            name,
            column_counts.get(name, ""),
            bool_label(first_value(meta, ["IsHidden", "Hidden", "isHidden"]), labels),
            first_value(meta, ["Description", "description"]) or "",
        ])
    return rows


def build_column_rows(schema, tmschema_columns=None, tmschema_tables=None, labels=None):
    labels = labels or doc_text("pt-BR")
    """Build data dictionary rows merging schema + TMSCHEMA_COLUMNS metadata.

    tmschema_columns uses a numeric TableID foreign key, so we first build a
    TableID → TableName map from tmschema_tables, then index by (table_name, col_name).
    """
    # Build TableID -> TableName map from tmschema_tables
    table_id_to_name = {}
    for row in (tmschema_tables or []):
        tid = first_value(row, ["ID", "TableID", "id"])
        name = first_value(row, ["Name", "TableName", "name"])
        if tid and name:
            table_id_to_name[str(tid)] = str(name)

    # Build lookup (table_name_lower, col_name_lower) -> tmschema row
    tmschema_lookup = {}
    for row in (tmschema_columns or []):
        tid = first_value(row, ["TableID", "Table", "table"])
        col = first_value(row, ["ExplicitName", "Name", "ColumnName", "name"])
        if not col:
            continue
        # Resolve TableID to name if possible
        table_name = table_id_to_name.get(str(tid), "") if tid else ""
        if table_name and not is_system_table(table_name):
            tmschema_lookup[(table_name.lower(), str(col).lower())] = row

    # Normalize DataType codes from TMSCHEMA to readable names
    DATATYPE_MAP = labels["datatype_map"]

    rows = []
    seen = set()
    for row in schema:
        table = first_value(row, ["TableName", "Table", "table"])
        column = first_value(row, ["ColumnName", "Name", "Column", "column", "name"])
        if not table or not column or is_system_table(table):
            continue
        key = (str(table).lower(), str(column).lower())
        if key in seen:
            continue
        seen.add(key)

        # Enrich with tmschema_columns if we resolved the name
        tm = tmschema_lookup.get(key, {})

        raw_type = (
            first_value(tm, ["ExplicitDataType", "DataType", "dataType"])
            or first_value(row, ["DataType", "Type", "type"])
            or ""
        )
        data_type = DATATYPE_MAP.get(str(raw_type).lower(), str(raw_type)) if raw_type else ""

        fmt = (
            first_value(tm, ["FormatString", "Format", "format"])
            or first_value(row, ["FormatString", "Format", "format"])
            or ""
        )

        is_hidden_val = (
            first_value(tm, ["IsHidden", "Hidden", "isHidden"])
            or first_value(row, ["IsHidden", "Hidden", "isHidden"])
        )
        hidden = bool_label(is_hidden_val, labels)

        expression = (
            first_value(tm, ["Expression", "expression"])
            or first_value(row, ["Expression", "expression"])
            or ""
        )

        rows.append([table, column, data_type, fmt, hidden, mask_secrets(expression)])

    # NOTE: We do NOT add rows directly from tmschema_columns here because its
    # TableID field is a numeric foreign key (e.g. 6107, 15), not a table name.
    # Without a resolved tmschema_tables join, spurious numeric "table names"
    # would appear in the dictionary.
    return rows


def build_expression_rows(raw_records, node_records, kind):
    rows = []
    seen = set()
    for row in raw_records:
        name = first_value(row, ["Name", "MeasureName", "ColumnName", "Measure", "Column", "name"])
        table = first_value(row, ["TableName", "Table", "table"])
        expression = first_value(row, ["Expression", "expression"]) or ""
        if kind == "calc_column" and not expression:
            continue
        # Skip system/internal tables (DateTableTemplate_, LocalDateTable_, etc.)
        if table and is_system_table(table):
            continue
        key = (str(table).lower(), str(name).lower(), str(expression).lower())
        if not name or key in seen:
            continue
        seen.add(key)
        # G7: mascara so no momento de gravar a linha -- a dedupe key acima
        # usa a expressao crua para nao colidir medidas distintas cujo unico
        # segredo mascarado as tornaria "iguais".
        rows.append([table, name, mask_secrets(expression)])

    if rows:
        return rows

    for node in node_records:
        meta = node.get("meta", {})
        expression = meta.get("expression", "")
        if kind == "calc_column" and not expression:
            continue
        table = meta.get("table", "")
        # Skip system/internal tables from node_records too
        if table and is_system_table(table):
            continue
        rows.append([table, node.get("label", ""), mask_secrets(expression)])
    return rows


def set_table_width(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)


def set_table_borders(table, color="D8DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = margins.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def add_paragraph_bottom_border(paragraph, color="D8DEE8", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_font(cell, color=None, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def bold_cell(cell):
    set_cell_font(cell, bold=True)


def clip(value, limit=900, labels=None):
    labels = labels or doc_text("pt-BR")
    text = doc_value(value)
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 18)].rstrip() + f"\n...[{labels['truncated']}]"


def normalize_code_text(value):
    text = doc_value(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def summarize_m_expression(expression, labels=None):
    labels = labels or doc_text("pt-BR")
    text = doc_value(expression)
    if not text.strip():
        return ""
    steps = re.findall(r"^\s*#?\"?([^\"=\r\n]{2,80})\"?\s*=", text, flags=re.MULTILINE)
    steps = [step.strip() for step in steps if step.strip().lower() not in {"let", "in"}]
    connector_match = re.search(r"([A-Za-z][A-Za-z0-9_]*\.[A-Za-z][A-Za-z0-9_]*)\s*\(", text)
    parts = []
    if connector_match:
        parts.append(f"{labels['connector']}: {connector_match.group(1)}")
    if steps:
        parts.append(f"{labels['steps']}: " + ", ".join(steps[:6]))
    if not parts:
        parts.append(clip(text, 180, labels))
    return " | ".join(parts)


def summarize_expression(expression, labels=None):
    labels = labels or doc_text("pt-BR")
    text = re.sub(r"\s+", " ", doc_value(expression)).strip()
    if not text:
        return ""
    functions = sorted(set(re.findall(r"\b([A-Z][A-Z0-9_]+)\s*\(", text)))
    if functions:
        return f"{labels['functions']}: " + ", ".join(functions[:8])
    return clip(text, 180, labels)


def bool_label(value, labels=None):
    labels = labels or doc_text("pt-BR")
    if value in ("", None):
        return ""
    return labels["yes"] if str(value).strip().lower() in {"true", "1", "yes", "sim"} else labels["no"]


def unique_texts(values):
    result = []
    seen = set()
    for value in values:
        text = str(value)
        if text not in seen:
            seen.add(text)
            result.append(text)
    return result
