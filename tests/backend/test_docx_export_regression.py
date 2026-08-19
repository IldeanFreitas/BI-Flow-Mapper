"""Regressao do refactor G20: build_documentation_docx() reaproveita os 8
helpers build_*_rows extraidos e usados tambem por build_documentation_html
-- este teste garante que o DOCX gerado continua um documento .docx VALIDO
depois da extracao (nao so "o servidor nao caiu"), abrindo a resposta com
python-docx de verdade e lendo document.xml.

Mesmo padrao de servidor real (live_server) + PBIXRay fake (_pbix_fixtures)
usado em test_export_html.py, para nao depender de um .pbix binario real.
"""
from __future__ import annotations

import io

import requests
from docx import Document

import backend

from _pbix_fixtures import make_empty_pbix_bytes, patch_pbixray

MALICIOUS_TABLE_NAME = 'Sales & "Special" <script>alert(1)</script>'


def _records():
    return {
        "tmschema_tables": [{"Name": MALICIOUS_TABLE_NAME, "IsHidden": False}],
        "schema": [{"TableName": MALICIOUS_TABLE_NAME, "ColumnName": "Amount", "DataType": "Int64"}],
        "dax_measures": [
            {"TableName": MALICIOUS_TABLE_NAME, "Name": "Total Amount", "Expression": "SUM('Sales'[Amount])"}
        ],
        "tmschema_datasources": [
            {"Name": "SalesDb", "Kind": "Sql", "ConnectionString": "Data Source=srv;Pwd=TopSecret123;"}
        ],
    }


def test_export_docx_endpoint_still_returns_a_valid_docx(live_server, monkeypatch):
    patch_pbixray(monkeypatch, backend, _records())

    response = requests.post(
        f"{live_server}/api/export-docx",
        data=make_empty_pbix_bytes(),
        headers={"Origin": live_server, "Content-Type": "application/octet-stream"},
        timeout=15,
    )
    assert response.status_code == 200, response.text
    assert (
        response.headers["Content-Type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.headers["Content-Disposition"].startswith("attachment;")

    # Abre com python-docx de verdade -- se o refactor tivesse corrompido a
    # estrutura OOXML (ex. uma tabela mal fechada, um helper devolvendo um
    # tipo que quebra add_records_table), isso levantaria excecao aqui.
    document = Document(io.BytesIO(response.content))

    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "BI Flow Mapper" in full_text

    # A mesma linha mascarada (G7) que apareceria no HTML tambem deve
    # aparecer no DOCX -- os helpers build_*_rows sao os MESMOS reaproveitados
    # pelos dois formatos.
    table_texts = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    joined_table_text = "\n".join(table_texts)
    assert "TopSecret123" not in joined_table_text
    assert "***MASKED***" in joined_table_text

    # document.xml em si precisa ser XML bem-formado (prova adicional, alem
    # de Document() nao ter levantado excecao, de que nada corrompeu a
    # arvore OOXML).
    document_xml = document.element.xml
    assert document_xml.startswith("<?xml") or "<w:document" in document_xml
    assert len(document.tables) > 0
